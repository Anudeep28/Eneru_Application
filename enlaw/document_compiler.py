import os
from docx import Document
from docx.shared import Inches
from django.template.loader import render_to_string
from django.conf import settings
import markdown
import re

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

class DocumentCompiler:
    def __init__(self, document):
        self.document = document
        self.template = document.template
        self.user_inputs = document.user_inputs
        
    def _replace_placeholders(self):
        """Replace placeholders in template content with user inputs"""
        content = self.template.content
        for key, value in self.user_inputs.items():
            placeholder = f"{{{{{key}}}}}"
            content = content.replace(placeholder, value)
        return content
    
    def to_pdf(self, output_path=None):
        """Convert document to PDF"""
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("WeasyPrint is not available. Please install it to generate PDFs.")
            
        try:
            # Convert markdown to HTML
            content = self._replace_placeholders()
            html_content = markdown.markdown(content)
            
            # Create HTML with proper styling
            html_string = render_to_string('enlaw/export_template.html', {
                'content': html_content,
                'title': self.template.title
            })
            
            # Generate PDF
            if not output_path:
                output_path = os.path.join(
                    settings.MEDIA_ROOT,
                    'documents',
                    f'{self.document.id}.pdf'
                )
                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            HTML(string=html_string).write_pdf(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error generating PDF: {str(e)}")
    
    def to_docx(self, output_path=None):
        """Convert document to DOCX"""
        try:
            # Create new Document
            doc = Document()
            
            # Add title
            doc.add_heading(self.template.title, 0)
            
            # Replace placeholders and convert content
            content = self._replace_placeholders()
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                # Handle markdown headings
                if para.startswith('#'):
                    level = len(re.match(r'^#+', para).group())
                    text = para.lstrip('#').strip()
                    doc.add_heading(text, level)
                    continue
                
                # Handle markdown lists
                if para.startswith('- ') or para.startswith('* '):
                    for line in para.split('\n'):
                        if line.strip():
                            doc.add_paragraph(
                                line.lstrip('- ').lstrip('* '),
                                style='List Bullet'
                            )
                    continue
                
                # Handle numbered lists
                if re.match(r'^\d+\.', para):
                    for line in para.split('\n'):
                        if line.strip():
                            doc.add_paragraph(
                                re.sub(r'^\d+\.\s*', '', line),
                                style='List Number'
                            )
                    continue
                
                # Regular paragraph
                doc.add_paragraph(para)
            
            # Save document
            if not output_path:
                output_path = os.path.join(
                    settings.MEDIA_ROOT,
                    'documents',
                    f'{self.document.id}.docx'
                )
                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error generating DOCX: {str(e)}")
