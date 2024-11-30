import sys
import os
import time
import wave
import pyaudio
import numpy as np
import whisper
import threading
import traceback
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QTextEdit, QVBoxLayout, QHBoxLayout, QWidget
from PyQt5.QtCore import pyqtSignal, QObject
from docx import Document
import lmproof as lm
from queue import Queue

proof = lm.load("en")

def Proofread(text, proof):
    error_free_text = proof.proofread(text)
    return error_free_text

class AudioRecorder(QObject):
    def __init__(self, buffer_queue):
        super().__init__()
        self.format = pyaudio.paFloat32
        self.channels = 1
        self.rate = 16000
        self.chunk = 1024
        self.recording = False
        self.paused = False
        self.audio = pyaudio.PyAudio()
        self.stream = None
        self.buffer_queue = buffer_queue

    def start_recording(self):
        self.recording = True
        self.paused = False
        self.stream = self.audio.open(format=self.format, channels=self.channels,
                                      rate=self.rate, input=True,
                                      frames_per_buffer=self.chunk)
        threading.Thread(target=self._record).start()

    def stop_recording(self):
        self.recording = False
        if self.stream:
            self.stream.stop_stream()
            self.stream.close()

    def pause_recording(self):
        self.paused = not self.paused

    def _record(self):
        while self.recording:
            if not self.paused:
                try:
                    data = self.stream.read(self.chunk)
                    audio_chunk = np.frombuffer(data, dtype=np.float32)
                    self.buffer_queue.put(audio_chunk)
                except Exception as e:
                    print(f"Error recording audio: {e}")
            else:
                time.sleep(0.1)

class AudioProcessor(QObject):
    chunk_ready = pyqtSignal(np.ndarray, float)

    def __init__(self, buffer_queue):
        super().__init__()
        self.buffer_queue = buffer_queue
        self.processing = False
        self.silence_threshold = 0.01
        self.min_chunk_duration = 5
        self.max_chunk_duration = 30

    def start_processing(self):
        self.processing = True
        threading.Thread(target=self._process).start()

    def stop_processing(self):
        self.processing = False

    def _process(self):
        frames = []
        start_time = time.time()
        last_sound_time = start_time

        while self.processing:
            if not self.buffer_queue.empty():
                audio_chunk = self.buffer_queue.get()
                frames.append(audio_chunk)

                if np.abs(audio_chunk).mean() > self.silence_threshold:
                    last_sound_time = time.time()

                chunk_duration = time.time() - start_time
                silence_duration = time.time() - last_sound_time

                if (chunk_duration >= self.min_chunk_duration and silence_duration >= 0.5) or \
                   chunk_duration >= self.max_chunk_duration:
                    audio_data = np.concatenate(frames)
                    self.chunk_ready.emit(audio_data, chunk_duration)
                    frames = []
                    start_time = time.time()
                    last_sound_time = start_time
            else:
                time.sleep(0.01)

class TranscriptionApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Whisper Realtime Transcription")
        self.setGeometry(100, 100, 600, 400)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout(self.central_widget)

        self.text_area = QTextEdit()
        self.layout.addWidget(self.text_area)

        self.button_layout = QHBoxLayout()
        self.start_button = QPushButton("Start Recording")
        self.stop_button = QPushButton("Stop Recording")
        self.pause_button = QPushButton("Pause/Resume")
        self.save_button = QPushButton("Save Transcript")

        self.button_layout.addWidget(self.start_button)
        self.button_layout.addWidget(self.stop_button)
        self.button_layout.addWidget(self.pause_button)
        self.button_layout.addWidget(self.save_button)

        self.layout.addLayout(self.button_layout)

        self.start_button.clicked.connect(self.start_recording)
        self.stop_button.clicked.connect(self.stop_recording)
        self.pause_button.clicked.connect(self.pause_recording)
        self.save_button.clicked.connect(self.save_transcript)

        self.buffer_queue = Queue()
        self.recorder = AudioRecorder(self.buffer_queue)
        self.processor = AudioProcessor(self.buffer_queue)
        self.processor.chunk_ready.connect(self.process_audio)

        try:
            model_path = r"C:\Users\PC\Documents\Speech2Text\Custom_Whisper_fine_tuning\whisper_models\whisper_small_version_40_s_1.pth"
            self.model = whisper.load_model(model_path)
            print("Whisper model loaded successfully")
        except Exception as e:
            print(f"Error loading Whisper model: {e}")
            self.model = None

        self.transcribed_text = ""
        self.prompt = "Transcribe the following Indian english audio accurately:"
        self.silence_threshold = 0.01  # Adjust this value as needed
        self.accumulated_audio = np.array([], dtype=np.float32)
        self.accumulated_duration = 0
        self.max_accumulated_duration = 30  # Maximum duration in seconds before forced transcription

    def start_recording(self):
        if self.model is None:
            self.text_area.setPlainText("Error: Whisper model not loaded.")
            return
        self.recorder.start_recording()
        self.processor.start_processing()
        self.start_button.setEnabled(False)
        self.stop_button.setEnabled(True)
        self.pause_button.setEnabled(True)
        self.text_area.setPlainText("Recording started...")

    def stop_recording(self):
        self.recorder.stop_recording()
        self.processor.stop_processing()
        self.start_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        self.pause_button.setEnabled(False)
        self.text_area.append("\nRecording stopped.")

    def pause_recording(self):
        self.recorder.pause_recording()
        status = "paused" if self.recorder.paused else "resumed"
        self.text_area.append(f"\nRecording {status}.")

    def process_audio(self, audio_data, duration):
        if self.model is None:
            self.text_area.setPlainText("Error: Whisper model not loaded.")
            return

        try:
            self.text_area.append(f"\nProcessing audio chunk of duration: {duration:.2f} seconds")

            # Accumulate audio data
            self.accumulated_audio = np.concatenate((self.accumulated_audio, audio_data))
            self.accumulated_duration += duration

            # Check if the new chunk contains speech
            if np.abs(audio_data).mean() <= self.silence_threshold:
                self.text_area.append("\nSilence detected in the new chunk.")
            else:
                self.text_area.append("\nSpeech detected in the new chunk.")

            # Decide whether to transcribe based on accumulated duration or max duration reached
            if self.accumulated_duration >= self.max_accumulated_duration or np.abs(audio_data).mean() > self.silence_threshold:
                self.text_area.append(f"\nTranscribing accumulated audio of duration: {self.accumulated_duration:.2f} seconds")

                # Pad or trim the audio to 30 seconds
                target_length = 30 * self.recorder.rate
                if len(self.accumulated_audio) < target_length:
                    padded_audio = np.pad(self.accumulated_audio, (0, target_length - len(self.accumulated_audio)))
                else:
                    padded_audio = self.accumulated_audio[:target_length]

                self.text_area.append("\nStarting transcription...")
                result = self.model.transcribe(padded_audio, language="en", initial_prompt=self.prompt)
                transcribed_chunk = result["text"].strip()
                transcribed_chunk = Proofread(transcribed_chunk, proof)
                
                if transcribed_chunk:
                    self.transcribed_text += transcribed_chunk + " "
                    self.text_area.setPlainText(self.transcribed_text)
                    self.text_area.append("\nTranscription successful.")
                    # Update the prompt with the last transcribed text
                    self.prompt = f"Transcribe the following audio accurately. Previous transcription: {transcribed_chunk}"
                else:
                    self.text_area.append("\nNo speech detected in the accumulated audio.")

                # Reset accumulated audio and duration
                self.accumulated_audio = np.array([], dtype=np.float32)
                self.accumulated_duration = 0
            else:
                self.text_area.append("\nAccumulating more audio before transcription.")

        except Exception as e:
            self.text_area.append(f"\nError in transcription: {str(e)}")
            self.text_area.append(f"\nError details: {type(e).__name__}")
            self.text_area.append(f"\nTraceback: {traceback.format_exc()}")

    def save_transcript(self):
        try:
            # Save as Markdown
            with open("transcript.md", "w") as f:
                f.write(self.transcribed_text)

            # Save as Word document
            doc = Document()
            doc.add_heading('Transcription of Speech to Text', level=1)  # Add a title
            doc.add_paragraph(self.transcribed_text)
            doc.save("transcript.docx")

            self.text_area.append("\nTranscript saved as transcript.md and transcript.docx")
        except Exception as e:
            self.text_area.append(f"\nError saving transcript: {e}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = TranscriptionApp()
    window.show()
    sys.exit(app.exec_())
